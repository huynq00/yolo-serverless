#!/usr/bin/env python3
"""Generate 20-minute presentation script (Word) for Thuyet_Trinh_YOLO_Serverless.pptx."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Script_Thuyet_Trinh_YOLO_Serverless.docx"

# Timing budget (minutes) — total ≈ 20
# Opening 1.5 | BG 2.5 | Goals 1.5 | Tech 2.5 | Arch 2.0 |
# Impl 1.5 | Test 1.5 | Results 5.0 | Close 2.0


def set_run_font(run, *, size=11, bold=False, italic=False, color=None, name="Times New Roman"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = color


def add_heading_custom(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    sizes = {0: 18, 1: 14, 2: 12, 3: 11}
    set_run_font(run, size=sizes.get(level, 11), bold=True,
                 color=RGBColor(0x1E, 0x3A, 0x8A))
    return p


def add_body(doc, text, *, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("💡 Giải thích: ")
    set_run_font(run, size=10, bold=True, color=RGBColor(0x06, 0xB6, 0xD4))
    run2 = p.add_run(text)
    set_run_font(run2, size=10, italic=True, color=RGBColor(0x33, 0x41, 0x55))
    return p


def add_tip(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("⏱ Gợi ý trình bày: ")
    set_run_font(run, size=10, bold=True, color=RGBColor(0x16, 0xA3, 0x4A))
    run2 = p.add_run(text)
    set_run_font(run2, size=10, color=RGBColor(0x33, 0x41, 0x55))
    return p


def add_slide_block(doc, slide_no, title, time_sec, script, explain=None, tip=None):
    add_heading_custom(doc, f"Slide {slide_no} — {title}  [{time_sec}s]", level=2)
    add_body(doc, "Lời thoại:", bold=True, size=11)
    # Allow multi-paragraph script
    for para in script.strip().split("\n\n"):
        add_body(doc, para.strip())
    if explain:
        add_note(doc, explain)
    if tip:
        add_tip(doc, tip)


def build():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # ===== COVER =====
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("SCRIPT THUYẾT TRÌNH ĐỒ ÁN")
    set_run_font(r, size=18, bold=True, color=RGBColor(0x0F, 0x17, 0x2A))

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run(
        "Tối ưu hóa thời gian khởi động lạnh cho suy diễn AI\n"
        "trên nền tảng Serverless bằng chia sẻ trọng số"
    )
    set_run_font(r2, size=13, bold=True, color=RGBColor(0x1E, 0x3A, 0x8A))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(
        "Môn: Hệ tính toán phân bố nâng cao — NT2204.CH201\n"
        "Học viên: Ngô Quang Huy  ·  GVHD: TS. Huỳnh Văn Đặng\n"
        "File slide: Thuyet_Trinh_YOLO_Serverless.pptx  ·  Thời lượng: ~20 phút"
    )
    set_run_font(rm, size=11, color=RGBColor(0x33, 0x41, 0x55))

    # ===== HOW TO USE =====
    add_heading_custom(doc, "Cách dùng script này", level=1)
    add_body(
        doc,
        "Mỗi slide có: (1) Lời thoại — đọc/diễn đạt khi trình bày; "
        "(2) Giải thích — dành cho người nghe hoặc bản thân khi cần làm rõ thuật ngữ; "
        "(3) Gợi ý trình bày — cách chỉ slide, nhấn mạnh số liệu. "
        "Thời gian ghi bên cạnh tiêu đề slide; tổng khoảng 20 phút (không tính Q&A).",
    )
    add_body(
        doc,
        "Chiến lược 20 phút: Slide phân đoạn (số lớn xanh) chỉ nói 1 câu chuyển tiếp. "
        "Ưu tiên số liệu then chốt: −76.9%, 17.35s vs 75.00s, 0% lỗi, 1.38s model load. "
        "Nếu thiếu thời gian: rút gọn slide 18–19, 26, 28 — chỉ nêu kết luận.",
        italic=True,
    )

    # ===== TIMING TABLE =====
    add_heading_custom(doc, "Phân bổ thời gian tổng thể", level=1)
    table = doc.add_table(rows=10, cols=3)
    table.style = "Table Grid"
    headers = ["Phần", "Slide", "Thời gian"]
    data = [
        ("Mở đầu & mục lục", "1–3", "~1 phút 20 giây"),
        ("Bối cảnh & vấn đề", "4–6", "~2 phút 30 giây"),
        ("Mục tiêu & quy trình", "7–9", "~1 phút 30 giây"),
        ("Công nghệ & Weight-Sharing", "10–13", "~2 phút 30 giây"),
        ("Kiến trúc hệ thống", "14–16", "~2 phút"),
        ("Hiện thực hóa", "17–19", "~1 phút 30 giây"),
        ("Kịch bản thử nghiệm", "20–21", "~1 phút 30 giây"),
        ("Kết quả & thảo luận", "22–32", "~5 phút"),
        ("Kết luận & Q&A", "33–37", "~2 phút (+ Q&A)"),
    ]
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shading = cell._teProp if False else None
        from docx.oxml import OxmlElement
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1E3A8A")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run_font(run, size=10)

    doc.add_paragraph()

    # ===== GLOSSARY PREVIEW =====
    add_heading_custom(doc, "Thuật ngữ cần nắm trước khi thuyết trình", level=1)
    glossary_short = [
        ("Cold-start", "Lần đầu (hoặc sau khi tắt) hệ thống phải khởi tạo lại container — chậm hơn nhiều so với khi đang chạy sẵn."),
        ("Serverless", "Mô hình hạ tầng tự co giãn: không có request thì thu về 0 (scale-to-zero), có request thì tự tạo lại."),
        ("Weights / Trọng số", "File số học của mô hình AI đã huấn luyện (ở đây ~90MB). Phải nạp vào RAM trước khi suy diễn."),
        ("P99 Latency", "Trong 100 request, 99 request nhanh hơn mức này; phản ánh trải nghiệm xấu nhất gần như chắc chắn gặp."),
        ("tmpfs", "Hệ thống tệp nằm trên RAM — đọc/ghi rất nhanh, nhưng mất dữ liệu khi tắt máy."),
        ("Weight-Sharing", "Nhiều Pod dùng chung một bản weights trên node, không mỗi Pod tự tải một bản từ đĩa."),
        ("Baseline vs Optimized", "Baseline = đọc weights từ đĩa (đối chứng). Optimized = đọc từ RAM tmpfs (giải pháp đề xuất)."),
        ("Pod", "Đơn vị chạy nhỏ nhất trên Kubernetes — chứa container ứng dụng (FastAPI + YOLO)."),
        ("Knative", "Lớp Serverless trên Kubernetes: tự scale, scale-to-zero, quản lý request khi Pod chưa sẵn sàng."),
        ("Burst / Warm", "Burst = tải đột biến nhiều user ảo. Warm = Pod đã sẵn sàng, request đều đặn — độ trễ thấp nhất."),
    ]
    for term, meaning in glossary_short:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"• {term}: ")
        set_run_font(r1, size=10, bold=True)
        r2 = p.add_run(meaning)
        set_run_font(r2, size=10)

    # ========== SCRIPTS ==========
    add_heading_custom(doc, "SCRIPT THEO TỪNG SLIDE", level=0)

    # --- Opening ---
    add_heading_custom(doc, "Phần A — Mở đầu (≈ 1 phút 20 giây)", level=1)

    add_slide_block(
        doc, 1, "Trang bìa", 40,
        """Kính chào Thầy/Cô và các bạn.

Em là Ngô Quang Huy. Hôm nay em xin trình bày đồ án môn Hệ tính toán phân bố nâng cao với đề tài: “Tối ưu hóa thời gian khởi động lạnh cho suy diễn AI trên nền tảng Serverless bằng chia sẻ trọng số”.

Giảng viên hướng dẫn là Thầy TS. Huỳnh Văn Đặng. Em xin phép bắt đầu.""",
        explain="Cold-start = “khởi động lạnh”: hệ thống đang tắt (0 instance), có request đến thì phải dựng lại container và nạp mô hình — rất chậm. "
                "Serverless = hạ tầng tự mở/đóng theo nhu cầu, tiết kiệm chi phí khi không có người dùng. "
                "Weight-Sharing = chia sẻ file trọng số mô hình giữa nhiều tiến trình/Pod.",
        tip="Đứng thẳng, nhìn khán giả; nói rõ họ tên và tên đề tài một lần, không đọc dài dòng phụ đề.",
    )

    add_slide_block(
        doc, 2, "Phân đoạn: Nội dung trình bày", 15,
        """Phần mở đầu, em xin nêu nhanh lộ trình thuyết trình.""",
        tip="Slide phân đoạn chỉ chuyển tiếp — không dừng lâu.",
    )

    add_slide_block(
        doc, 3, "Mục lục 8 phần", 25,
        """Nội dung gồm tám phần chính: từ bối cảnh và vấn đề cold-start, mục tiêu đề tài, công nghệ và phương pháp Weight-Sharing, kiến trúc hệ thống, hiện thực hóa, các kịch bản thử nghiệm, kết quả định lượng, và cuối cùng là kết luận cùng hướng phát triển.

Trong khoảng 20 phút, em sẽ tập trung vào vấn đề, giải pháp, và các số liệu đối chứng quan trọng nhất.""",
        tip="Chỉ tay theo 3 mốc: Vấn đề → Giải pháp → Kết quả (−76.9%).",
    )

    # --- Background ---
    add_heading_custom(doc, "Phần B — Bối cảnh & vấn đề (≈ 2 phút 30 giây)", level=1)

    add_slide_block(
        doc, 4, "Phân đoạn: Bối cảnh & vấn đề", 10,
        """Em xin đi vào bối cảnh: vì sao AI kết hợp Serverless lại gặp bài toán cold-start.""",
    )

    add_slide_block(
        doc, 5, "Bối cảnh đề tài", 75,
        """Hai xu hướng đang gặp nhau. Một bên là Deep Learning — cần hạ tầng linh hoạt để chạy mô hình nhận diện. Bên kia là Serverless — tự co giãn và có thể thu về zero khi không có tải, giúp tiết kiệm tài nguyên.

Nhưng khi hai xu hướng này kết hợp, xuất hiện nút thắt lớn: mỗi lần hệ thống “ngủ” rồi thức dậy, container mới phải nạp lại file trọng số mô hình. Với YOLO-World khoảng 90 megabyte, nếu đọc từ đĩa thì rất chậm.

Hệ quả là độ trễ P99 tăng mạnh — nghĩa là nhóm request chậm nhất bị “đánh” rất nặng — ảnh hưởng trải nghiệm thời gian thực. Do đó đồ án hướng tới tái thiết kế lớp lưu trữ dùng chung trên cùng một node.""",
        explain="Deep Learning: học sâu, dùng mạng nơ-ron lớn. Scale-to-zero: không có request thì tắt hết Pod (0 bản sao). "
                "I/O (Input/Output): thao tác đọc/ghi dữ liệu; đọc file lớn từ đĩa chậm hơn RAM rất nhiều. "
                "Container: “hộp” chạy ứng dụng độc lập. Node: một máy (ảo) trong cụm. "
                "P99: nếu 100 người gọi API, 99 người nhanh hơn mức P99 — mức này phản ánh “xấu nhất gần như chắc chắn”.",
        tip="Nói chậm ở câu “đọc từ đĩa thì rất chậm” — đây là động lực của đề tài.",
    )

    add_slide_block(
        doc, 6, "Vấn đề cần giải quyết", 65,
        """Tóm lại, vấn đề cốt lõi là: mỗi Pod mới phải đọc weights từ lưu trữ trên node, và I/O đĩa trở thành nút thắt.

Câu hỏi nghiên cứu của đồ án: làm sao giảm chi phí nạp weights mà vẫn giữ được kiến trúc Serverless — tức vẫn scale-to-zero được?

Hướng tiếp cận em đề xuất là Weight-Sharing qua phân vùng tmpfs trên RAM, gắn vào Pod bằng hostPath. Để chứng minh hiệu quả, em đối chứng với Baseline: cùng mô hình nhưng đọc từ đĩa ảo trên Minikube.""",
        explain="Pod: đơn vị chạy nhỏ nhất trên Kubernetes. "
                "tmpfs: “ổ đĩa giả” nằm trên RAM — nhanh, nhưng dữ liệu mất khi restart. "
                "hostPath: cơ chế Kubernetes gắn thư mục của máy node vào trong Pod. "
                "Baseline: nhóm đối chứng (cách làm thông thường). Optimized: nhóm tối ưu (giải pháp đề xuất). "
                "Minikube: cụm Kubernetes giả lập trên một máy, phù hợp thí nghiệm đồ án.",
        tip="Nhấn mạnh câu hỏi nghiên cứu — giúp hội đồng nắm “em giải quyết cái gì”.",
    )

    # --- Goals ---
    add_heading_custom(doc, "Phần C — Mục tiêu & quy trình (≈ 1 phút 30 giây)", level=1)

    add_slide_block(
        doc, 7, "Phân đoạn: Mục tiêu", 10,
        """Mục tiêu cốt lõi: giảm đáng kể độ trễ cold-start khi nạp weights mô hình AI trên Serverless.""",
    )

    add_slide_block(
        doc, 8, "Mục tiêu cụ thể", 50,
        """Để đạt mục tiêu đó, đồ án làm ba việc chính.

Thứ nhất, dựng hạ tầng Serverless chuẩn: Minikube cộng Knative Serving, có autoscaling và scale-to-zero.

Thứ hai, hiện thực Weight-Sharing: cấp phát tmpfs 4GB trên RAM, các Pod dùng chung qua hostPath.

Thứ ba, đo lường đối chứng bằng k6, Prometheus và Grafana — so sánh RAM với Disk theo P99, Burst và Warm.""",
        explain="Autoscaling: tự tăng/giảm số Pod theo tải. "
                "k6: công cụ giả lập nhiều người dùng ảo gửi request. "
                "Prometheus: thu thập số liệu nội bộ (metrics). Grafana: vẽ biểu đồ từ Prometheus. "
                "Burst: tải đột biến. Warm: hệ thống đã “nóng”, Pod sẵn sàng.",
    )

    add_slide_block(
        doc, 9, "Quy trình 4 giai đoạn", 30,
        """Quy trình gồm bốn giai đoạn: khởi tạo hạ tầng; cấu hình Weight-Sharing; chạy load test các kịch bản Cold, Burst, Warm; rồi tổng hợp số liệu và đánh giá P99.

Em xin chuyển sang phần công nghệ và phương pháp.""",
        tip="Chỉ lần lượt 4 khối trên hình, mỗi khối một câu ngắn.",
    )

    # --- Tech ---
    add_heading_custom(doc, "Phần D — Công nghệ & Weight-Sharing (≈ 2 phút 30 giây)", level=1)

    add_slide_block(
        doc, 10, "Phân đoạn: Công nghệ", 10,
        """Phần này trình bày stack công nghệ và ý tưởng Weight-Sharing qua tmpfs.""",
    )

    add_slide_block(
        doc, 11, "Stack công nghệ", 55,
        """Về hạ tầng: Kubernetes trên Minikube, Knative Serving với Autoscaler và Activator, cổng Kourier, và volume hostPath để chia sẻ weights.

Về ứng dụng và giám sát: mô hình YOLO-World của Ultralytics, API FastAPI, Prometheus với Grafana, và k6 để tạo tải cold, burst và chu kỳ đầy đủ.

Toàn bộ thí nghiệm chạy trên một node để đảm bảo tính cục bộ của dữ liệu khi so sánh RAM và Disk.""",
        explain="Kubernetes: hệ điều phối container. Knative Serving: biến K8s thành Serverless. "
                "Activator: “người gác cổng” giữ request khi chưa có Pod, rồi kích hoạt scale-from-zero. "
                "Kourier: cổng vào (Ingress) đưa HTTP vào cụm. "
                "YOLO-World: mô hình nhận diện đối tượng trong ảnh. FastAPI: framework viết API Python nhanh. "
                "End-to-end: đo từ lúc client gửi đến lúc nhận đủ phản hồi.",
    )

    add_slide_block(
        doc, 12, "Optimized vs Baseline", 50,
        """Hai cơ chế đối chứng rất rõ.

Bên Optimized: weights nằm trên tmpfs RAM tại /mnt/shared-weights. Các Pod đọc chung một bản — P99 cold trung bình khoảng 17,35 giây.

Bên Baseline: cùng file nhưng trên đĩa tại /mnt/disk-weights — mỗi lần cold-start phải đọc đĩa, P99 trung bình khoảng 75 giây.

Đây chính là “phép thử” để đo mức cải thiện do Weight-Sharing.""",
        explain="Cùng workload, cùng node, khác duy nhất nơi đặt file weights → kết luận về I/O đáng tin hơn. "
                "17s vs 75s: khoảng chênh hơn 4 lần — đây là con số cần thuộc lòng.",
        tip="Dùng tay trái/phải chỉ hai cột Optimized / Baseline trên slide.",
    )

    add_slide_block(
        doc, 13, "Hình Weight-Sharing", 35,
        """Hình này minh họa: nhiều Pod trên cùng node cùng trỏ về một thư mục hostPath. Bên trái là RAM tmpfs — nhanh; bên phải là Disk — chậm hơn rõ rệt.

Ý tưởng then chốt: không nhân bản I/O đĩa cho mỗi lần khởi động Pod.""",
        tip="Dừng 3–5 giây để khán giả nhìn hình; nói ít, để hình “làm việc”.",
    )

    # --- Architecture ---
    add_heading_custom(doc, "Phần E — Kiến trúc (≈ 2 phút)", level=1)

    add_slide_block(
        doc, 14, "Phân đoạn: Kiến trúc", 10,
        """Tiếp theo là kiến trúc tổng thể của pipeline suy diễn.""",
    )

    add_slide_block(
        doc, 15, "Kiến trúc tổng thể", 55,
        """Luồng xử lý như sau: k6 gửi ảnh lên endpoint /predict; Kourier nhận và định tuyến; Knative Serving điều phối autoscaling — gồm Activator và Autoscaler; request vào Pod chạy FastAPI cộng YOLO-World.

Song song, Prometheus cạo metric từ /metrics và Grafana hiển thị. Dưới cùng là lớp lưu trữ dùng chung: tmpfs cho Optimized và disk cho Baseline.""",
        explain="Endpoint /predict: địa chỉ API nhận ảnh và trả kết quả nhận diện. "
                "Định tuyến: chọn đúng dịch vụ dựa trên tên miền ảo. "
                "Queue-Proxy: sidecar đi kèm Pod, giới hạn số request đồng thời. "
                "Scrape: Prometheus định kỳ “kéo” số liệu từ ứng dụng.",
        tip="Theo mũi tên trên hình từ trái sang phải: k6 → Kourier → Knative → Pod.",
    )

    add_slide_block(
        doc, 16, "Các thành phần chính", 55,
        """Em xin nhấn sáu thành phần: k6 đo độ trễ end-to-end; Kourier là cổng vào; Knative lo scale-from-zero; Pod chứa logic suy diễn và nạp model ngay khi process khởi động; lớp hostPath quyết định tốc độ nạp weights; Prometheus/Grafana quan sát nội tại.

Điểm thiết kế quan trọng: nạp model ở top-level — nghĩa là chi phí I/O xảy ra trước khi Pod báo Ready.""",
        explain="Ready: trạng thái Kubernetes báo Pod đã sẵn sàng nhận traffic. "
                "Top-level import: nạp model lúc module Python load, không đợi request đầu tiên trong handler — "
                "giúp đo cold-start “thật” vì Pod chưa Ready cho đến khi model vào RAM xong.",
    )

    # --- Implementation ---
    add_heading_custom(doc, "Phần F — Hiện thực hóa (≈ 1 phút 30 giây)", level=1)

    add_slide_block(
        doc, 17, "Phân đoạn: Hiện thực", 8,
        """Phần hiện thực hóa tập trung vào cấu hình Knative và API FastAPI.""",
    )

    add_slide_block(
        doc, 18, "Cấu hình triển khai", 45,
        """Em triển khai hai dịch vụ Knative độc lập: optimized và baseline. Cấu hình autoscaling: min-scale bằng 0, max-scale bằng 5, target concurrency bằng 1.

imagePullPolicy đặt Never để dùng image cục bộ, loại trừ độ trễ kéo image từ mạng. Các script setup-weights và deploy-all giúp tái lập thí nghiệm ổn định.""",
        explain="min-scale=0: cho phép scale-to-zero. max-scale=5: tối đa 5 Pod — tránh “nổ” tài nguyên máy thí nghiệm. "
                "target concurrency=1: mỗi Pod xử lý khoảng 1 request đồng thời — phù hợp suy diễn nặng. "
                "imagePullPolicy Never: không tải image từ registry; dùng image đã build sẵn trên node.",
        tip="Nếu thiếu giờ: chỉ nói 2 dịch vụ đối chứng + min=0/max=5.",
    )

    add_slide_block(
        doc, 19, "API FastAPI", 37,
        """API có ba endpoint chính: /health kiểm tra trạng thái; /predict nhận ảnh và chạy YOLO; /metrics xuất số liệu Prometheus, trong đó có yolo_model_load_seconds — thời gian nạp model lúc startup.""",
        explain="JSON detections: kết quả nhận diện dạng cấu trúc (nhãn, tọa độ hộp…). "
                "Histogram latency: phân bố thời gian xử lý request, dùng để tính P99 phía ứng dụng.",
    )

    # --- Experiments ---
    add_heading_custom(doc, "Phần G — Kịch bản thử nghiệm (≈ 1 phút 30 giây)", level=1)

    add_slide_block(
        doc, 20, "Phân đoạn: Thử nghiệm", 10,
        """Chỉ số then chốt trong mọi kịch bản là P99 Latency.""",
        explain="Vì sao P99 chứ không phải trung bình? Trung bình dễ bị “làm đẹp” bởi nhiều request nhanh; "
                "P99 phản ánh nhóm chậm nhất — đúng với SLA và trải nghiệm người dùng khó tính.",
    )

    add_slide_block(
        doc, 21, "Bốn kịch bản", 80,
        """Kịch bản 1 — Cold-start: trước mỗi lần chạy ép scale-to-zero, gửi 3 user ảo đồng thời, lặp 3 lần cho cả RAM và Disk.

Kịch bản 2 — Burst: tăng dần lên 15 VU trong 15 giây, giữ 60 giây, rồi giảm. Tiêu chí nghiệm thu: lỗi HTTP dưới 10% và P99 dưới 60 giây.

Kịch bản 3 — Full cycle: Cold rồi Burst rồi Warm với 5 VU trong một phút.

Kịch bản 4 — Giám sát nội bộ metric thời gian nạp model qua Prometheus và Grafana.""",
        explain="VU (Virtual User): người dùng ảo do k6 giả lập. "
                "Ramp-up/ramp-down: tăng/giảm tải dần để mô phỏng thực tế. "
                "n=3: lặp 3 lần độc lập — mẫu nhỏ, mang tính minh họa ổn định tương đối, em sẽ nêu lại ở hạn chế.",
        tip="Ba chữ khóa: Cold – Burst – Warm. Hội đồng dễ nhớ.",
    )

    # --- Results ---
    add_heading_custom(doc, "Phần H — Kết quả & thảo luận (≈ 5 phút)", level=1)

    add_slide_block(
        doc, 22, "Phân đoạn: Kết quả", 12,
        """Đây là phần quan trọng nhất: Weight-Sharing qua RAM giảm trung bình 76,9% P99 cold-start so với đọc đĩa.""",
        tip="Nói chậm và rõ “bảy mươi sáu phẩy chín phần trăm”.",
    )

    add_slide_block(
        doc, 23, "Kết quả nổi bật", 55,
        """Con số headline: giảm 76,9% — từ 75 giây xuống 17,35 giây.

Optimized ổn định hơn với độ lệch chuẩn khoảng 1,68 giây; Baseline dao động mạnh hơn, khoảng 12,78 giây.

Trong Burst: không có lỗi HTTP, P99 khoảng 46,37 giây — dưới ngưỡng 60 giây. Thời gian nạp model từ RAM chỉ 1,38 giây. Khi Warm, P99 còn khoảng 5,64 giây.""",
        explain="σ (sigma): độ lệch chuẩn — đo mức “dao động” giữa các lần chạy. Optimized ít dao động hơn → ổn định hơn. "
                "0% lỗi HTTP: mọi request đều nhận được phản hồi thành công trong thí nghiệm burst. "
                "1,38s: chỉ là phần đọc/nạp file model; tổng cold-start còn gồm lập lịch Pod và khởi tạo Python.",
    )

    add_slide_block(
        doc, 24, "Bảng P99 cold-start", 40,
        """Bảng chi tiết ba lần chạy. Optimized dao động quanh 15 đến 19 giây; Baseline từ khoảng 60 đến 84 giây. Trung bình 17,35 so với 75 — khoảng cách rất rõ.""",
        tip="Chỉ hàng “Trung bình”; không cần đọc từng ô.",
    )

    add_slide_block(
        doc, 25, "Biểu đồ so sánh P99", 30,
        """Biểu đồ cột cho thấy cùng một kết luận trực quan: cột xanh Optimized thấp hơn rõ so với cột đỏ Baseline trên cả ba lần chạy và trung bình.""",
    )

    add_slide_block(
        doc, 26, "Đối chiếu dashboard style", 20,
        """Góc nhìn dạng dashboard: 17,35 giây đối 75 giây — mức cải thiện khoảng gần bốn lần.""",
        tip="Có thể bỏ qua nếu đã nói kỹ slide 23–25.",
    )

    add_slide_block(
        doc, 27, "Bảng Burst Traffic", 35,
        """Với Burst: 61 request, lỗi 0%, trung bình khoảng 19,4 giây, P99 khoảng 46,37 giây, không tắc nghẽn. Hệ thống đạt tiêu chí nghiệm thu đã đặt ra.""",
        explain="Congested = tắc nghẽn: hàng đợi request phình to, độ trễ tăng không kiểm soát. Thí nghiệm ghi nhận không tắc nghẽn.",
    )

    add_slide_block(
        doc, 28, "Biểu đồ phân vị Burst", 20,
        """Các phân vị Avg, P95, P99 đều nằm dưới ngưỡng 60 giây được vẽ bằng đường đứt nét đỏ.""",
        tip="Rút gọn nếu thiếu giờ.",
    )

    add_slide_block(
        doc, 29, "Bảng Cold–Burst–Warm", 35,
        """Chu kỳ đầy đủ: Cold P99 khoảng 22,4 giây; Burst khoảng 32,6 giây; Warm chỉ còn khoảng 5,6 giây với 39 request. Tổng 111 request thành công, lỗi 0%.""",
        explain="Warm thấp vì Pod đã sẵn sàng, model đã trong bộ nhớ — không còn trả giá cold-start. "
                "Burst cao hơn Warm vì vừa scale vừa xử lý tải lớn.",
    )

    add_slide_block(
        doc, 30, "Biểu đồ Cold–Burst–Warm", 25,
        """Hình này cho thấy rõ “hố” độ trễ lúc cold và mức ổn định lúc warm. Khoảng cách Cold trừ Warm giúp ước lượng chi phí khởi tạo hạ tầng mỗi lần scale-up.""",
    )

    add_slide_block(
        doc, 31, "Phân bổ chi phí cold-start", 40,
        """Phân tích sâu hơn: với Optimized, thời gian nạp model chỉ 1,38 giây — khoảng 8% của 17,35 giây P99. Phần còn lại chủ yếu là lập lịch container, khởi tạo runtime Python và deserialize.

Nghĩa là: sau khi đã tối ưu I/O đĩa, giới hạn tiếp theo nằm ở tầng hạ tầng và runtime — không còn là đọc file weights.""",
        explain="Deserialize: giải mã file weights thành cấu trúc tensor trong bộ nhớ để mô hình chạy được. "
                "Runtime Python: thời gian khởi động trình thông dịch và thư viện (Ultralytics, v.v.). "
                "Đây là insight quan trọng cho hướng phát triển.",
        tip="Nhấn câu “chỉ còn khoảng 8% là I/O model” — thể hiện em hiểu bottleneck đã dịch chuyển.",
    )

    add_slide_block(
        doc, 32, "Thảo luận", 40,
        """Tóm lại năm điểm: Weight-Sharing qua RAM xử lý hiệu quả nút thắt I/O; cold-start còn khoảng 17 giây chủ yếu do scheduling và runtime; Knative chịu burst tốt trong max-scale 5; pipeline đo lường có thể tái lập; và khoảng cách Cold–Warm hỗ trợ thiết lập SLA thực tế.""",
        explain="SLA (Service Level Agreement): cam kết chất lượng dịch vụ với người dùng, ví dụ “99% request dưới X giây”.",
    )

    # --- Conclusion ---
    add_heading_custom(doc, "Phần I — Kết luận & Q&A (≈ 2 phút)", level=1)

    add_slide_block(
        doc, 33, "Phân đoạn: Kết luận", 8,
        """Em xin chuyển sang kết luận, hạn chế và hướng phát triển.""",
    )

    add_slide_block(
        doc, 34, "Kết luận", 45,
        """Đồ án đã triển khai thành công kiến trúc AI Serverless trên Kubernetes và Knative. Weight-Sharing qua tmpfs giảm 76,9% P99 cold-start — từ 75 giây xuống 17,35 giây. Hệ thống ổn định dưới burst với 0% lỗi và P99 dưới 60 giây. Đồng thời xây dựng pipeline giám sát k6–Prometheus–Grafana.

Kết luận thực tiễn: đặt thành phần “nóng” — file weights — trên RAM là chiến lược hiệu quả cho Serverless AI.""",
    )

    add_slide_block(
        doc, 35, "Hạn chế & hướng phát triển", 50,
        """Về hạn chế: thí nghiệm trên Minikube một node; tmpfs 4GB là cấu hình tĩnh; cold-start Optimized vẫn còn khoảng 17 giây; mẫu n=3 mang tính minh họa.

Về hướng phát triển: cache phân tán kiểu Alluxio cho đa node; nén mô hình bằng quantization hoặc pruning; pre-warming thông minh theo dự báo tải; và chuyển sang engine suy diễn như Triton hoặc ONNX Runtime để giảm chi phí runtime Python.""",
        explain="Alluxio: lớp cache/lưu trữ phân tán, chia sẻ dữ liệu giữa nhiều node. "
                "Quantization: giảm độ chính xác số (ví dụ INT8) để file nhỏ hơn, chạy nhanh hơn. "
                "Pruning: cắt bớt tham số ít quan trọng trong mạng. "
                "Pre-warming: chủ động dựng Pod trước giờ cao điểm để tránh cold-start. "
                "Triton / ONNX Runtime: engine suy diễn tối ưu hơn stack Python thuần.",
        tip="Nêu hạn chế thành thật rồi nối ngay hướng phát triển — tạo ấn tượng khoa học tốt.",
    )

    add_slide_block(
        doc, 36, "Slide kết luận tổng", 20,
        """Tóm một câu: đồ án đã giảm đáng kể nút thắt I/O đĩa khi nạp trọng số AI trên Serverless; Weight-Sharing qua RAM hiệu quả, và bước tiếp theo là tối ưu scheduling cùng runtime.""",
    )

    add_slide_block(
        doc, 37, "Hỏi & Đáp", 15,
        """Em xin cảm ơn Thầy/Cô và các bạn đã lắng nghe. Em sẵn sàng trao đổi các câu hỏi.""",
        tip="Sau câu cảm ơn, dừng lại, mỉm cười — chờ câu hỏi. Không vội tắt slide.",
    )

    # ===== Q&A PREP =====
    add_heading_custom(doc, "Chuẩn bị trả lời câu hỏi thường gặp", level=1)

    faqs = [
        (
            "Vì sao dùng Minikube một node chứ không phải cụm thật?",
            "Mục tiêu đồ án là đối chứng I/O cục bộ RAM vs Disk trên cùng điều kiện. Single-node loại trừ biến số mạng giữa các node, phù hợp phạm vi môn học và tài nguyên máy cá nhân. Em đã nêu đây là hạn chế và đề xuất Alluxio/multi-node ở hướng phát triển.",
        ),
        (
            "n=3 có đủ thống kê không?",
            "Với mẫu nhỏ, P99 gần giá trị cực đại và mang tính minh họa mức cải thiện cùng độ ổn định tương đối. Em không tuyên bố suy luận thống kê chặt; xu hướng 3/3 lần Optimized đều thấp hơn Baseline rất rõ (≈15–19s vs ≈60–84s).",
        ),
        (
            "Sau khi tối ưu I/O, vì sao vẫn còn ~17 giây?",
            "Metric yolo_model_load_seconds ≈ 1,38s cho thấy phần đọc file chỉ ~8%. Phần còn lại thuộc lập lịch Pod, khởi tạo Python/Ultralytics và deserialize — ngoài phạm vi Weight-Sharing. Đây cũng là lý do hướng tới Triton/ONNX và pre-warming.",
        ),
        (
            "tmpfs mất dữ liệu khi restart thì có ổn không?",
            "Trong thiết kế Serverless, weights có thể được provision lại khi node khởi động (script setup-weights). tmpfs dùng cho lớp “nóng” tốc độ cao; nguồn bền vững vẫn có thể nằm trên đĩa hoặc registry và được đồng bộ vào RAM trước khi phục vụ.",
        ),
        (
            "Vì sao chọn YOLO-World?",
            "Đây là workload Computer Vision có file weights đủ lớn (~90MB) để lộ rõ nút thắt I/O, đồng thời là bài toán suy diễn thực tế (object detection), phù hợp minh họa Serverless AI.",
        ),
        (
            "Concurrency target = 1 có quá bảo thủ?",
            "Suy diễn YOLO trên CPU khá nặng; target=1 giúp đo hành vi scale rõ ràng và tránh quá tải một Pod. Có thể tinh chỉnh khi triển khai thực tế tùy SLA và tài nguyên GPU/CPU.",
        ),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        rq = p.add_run(f"Hỏi: {q}")
        set_run_font(rq, size=11, bold=True, color=RGBColor(0x1E, 0x3A, 0x8A))
        add_body(doc, f"Đáp: {a}")

    # ===== FULL GLOSSARY =====
    add_heading_custom(doc, "Phụ lục — Bảng giải thích thuật ngữ (tra cứu nhanh)", level=1)
    gloss = [
        ("AI / Deep Learning", "Trí tuệ nhân tạo / học sâu — dùng mạng nơ-ron nhiều lớp để nhận diện ảnh, giọng nói…"),
        ("Serverless", "Mô hình vận hành: nhà cung cấp tự quản lý server; ứng dụng scale theo request, có thể về 0."),
        ("Cold-start", "Khởi động từ trạng thái 0 instance — phải tạo Pod, nạp model → chậm."),
        ("Warm", "Pod đã chạy sẵn, model trong RAM → request nhanh."),
        ("Burst Traffic", "Lưu lượng tăng đột biến trong thời gian ngắn."),
        ("Scale-to-zero / Scale-from-zero", "Thu về 0 Pod khi rảnh / dựng lại Pod từ 0 khi có request."),
        ("Weights", "Tham số đã học của mô hình, lưu thành file (ví dụ .pt)."),
        ("Weight-Sharing", "Nhiều Pod dùng chung một bản weights trên node."),
        ("tmpfs", "Filesystem trên RAM — I/O nhanh, không bền vững qua reboot."),
        ("hostPath", "Gắn thư mục của node vào Pod trong Kubernetes."),
        ("Kubernetes (K8s)", "Hệ điều phối container trên cụm máy."),
        ("Minikube", "K8s giả lập một node trên máy local."),
        ("Knative Serving", "Lớp Serverless trên K8s (autoscaling, scale-to-zero)."),
        ("KPA / Activator / Queue-Proxy", "Autoscaler; thành phần giữ request khi scale-from-zero; sidecar giới hạn concurrency."),
        ("Kourier", "Ingress gateway của Knative."),
        ("Pod / Container", "Đơn vị chạy K8s / tiến trình đóng gói ứng dụng."),
        ("FastAPI", "Framework API Python hiệu năng cao."),
        ("YOLO-World", "Mô hình phát hiện đối tượng (object detection)."),
        ("Docker image", "Gói chứa ứng dụng + thư viện để chạy container."),
        ("k6", "Công cụ load test, giả lập VU gửi HTTP."),
        ("VU", "Virtual User — người dùng ảo."),
        ("Prometheus / Grafana", "Thu thập metrics / trực quan hóa dashboard."),
        ("P99 / Avg / P95", "Phân vị 99% / trung bình / phân vị 95% của độ trễ."),
        ("Latency", "Độ trễ — thời gian từ gửi request đến nhận response."),
        ("SLA", "Cam kết chất lượng dịch vụ với người dùng."),
        ("I/O", "Đọc/ghi dữ liệu (đĩa, mạng…)."),
        ("Baseline / Optimized", "Nhóm đối chứng (đĩa) / nhóm tối ưu (RAM)."),
        ("Deserialize", "Chuyển file weights thành cấu trúc trong bộ nhớ."),
        ("Quantization / Pruning", "Nén mô hình: giảm bit số / cắt tham số thừa."),
        ("Pre-warming", "Chủ động tạo Pod trước khi tải tăng."),
        ("Alluxio", "Lớp cache/lưu trữ phân tán đa node."),
        ("Triton / ONNX Runtime", "Engine suy diễn tối ưu thay cho Python thuần."),
    ]
    gtable = doc.add_table(rows=1 + len(gloss), cols=2)
    gtable.style = "Table Grid"
    from docx.oxml import OxmlElement
    for j, h in enumerate(["Thuật ngữ", "Giải thích dễ hiểu"]):
        cell = gtable.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1E3A8A")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for i, (term, meaning) in enumerate(gloss):
        gtable.rows[i + 1].cells[0].text = ""
        r1 = gtable.rows[i + 1].cells[0].paragraphs[0].add_run(term)
        set_run_font(r1, size=9, bold=True)
        gtable.rows[i + 1].cells[1].text = ""
        r2 = gtable.rows[i + 1].cells[1].paragraphs[0].add_run(meaning)
        set_run_font(r2, size=9)

    # ===== CHECKLIST =====
    add_heading_custom(doc, "Checklist trước giờ thuyết trình", level=1)
    checks = [
        "Thuộc 4 số: 76,9% · 17,35s · 75,00s · 0% lỗi / P99 burst 46,37s",
        "Thuộc insight: model load 1,38s ≈ 8% — bottleneck đã chuyển sang runtime",
        "Mở sẵn file Thuyet_Trinh_YOLO_Serverless.pptx, chế độ Presenter View nếu có",
        "Đồng hồ/điện thoại đặt chế độ không làm phiền; canh mốc phút 8 (xong kiến trúc) và phút 15 (xong kết quả chính)",
        "In hoặc mở song song file Word này trên máy phụ / điện thoại",
        "Chuẩn bị trả lời 2–3 câu hỏi ở mục FAQ phía trên",
    ]
    for c in checks:
        p = doc.add_paragraph()
        r = p.add_run(f"☐  {c}")
        set_run_font(r, size=11)

    add_body(
        doc,
        "\n— Hết script. Chúc bạn thuyết trình tốt!",
        italic=True,
        size=11,
    )

    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
