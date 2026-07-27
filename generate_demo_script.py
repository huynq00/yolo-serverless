#!/usr/bin/env python3
"""Generate video demo script (Word) — live k6 cold A/B, image v5 + MODEL_IO_MBPS."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Kich_Ban_Demo_YOLO_Serverless_5phut.docx"

URL_OPT = "http://yolo-inference.default.127.0.0.1.sslip.io:8080"
URL_BASE = "http://yolo-inference-baseline.default.127.0.0.1.sslip.io:8080"

# Số tham chiếu từ dry-run ĐẠT gần nhất (chỉ để biết kỳ vọng — khi quay đọc số LIVE).
REF_BASE_P99 = "59.49s"
REF_OPT_P99 = "18.07s"
REF_IMPROVE = "69.6%"


def set_run_font(run, *, size=11, bold=False, italic=False, color=None, name="Times New Roman"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = color


def add_heading_custom(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    sizes = {0: 18, 1: 14, 2: 12, 3: 11}
    set_run_font(run, size=sizes.get(level, 11), bold=True,
                 color=RGBColor(0x1E, 0x3A, 0x8A))
    return p


def add_body(doc, text, *, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_label(doc, label, text, color=RGBColor(0x16, 0xA3, 0x4A)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(label)
    set_run_font(run, size=10, bold=True, color=color)
    run2 = p.add_run(text)
    set_run_font(run2, size=10, color=RGBColor(0x33, 0x41, 0x55))
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text.strip())
    set_run_font(run, size=9, name="Consolas", color=RGBColor(0x1E, 0x29, 0x3B))
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p


def add_step(doc, step_no, where, action, command=None, expect=None, speak=None, tip=None, record=False):
    tag = " [QUAY]" if record else " [KHÔNG QUAY]"
    add_heading_custom(doc, f"Bước {step_no} — {action}{tag}", level=2)
    add_label(doc, "Nhập vào: ", where, color=RGBColor(0xDC, 0x26, 0x26))
    if command:
        add_body(doc, "Gõ / dán đúng lệnh sau rồi Enter:", bold=True, size=11)
        add_code(doc, command)
    if expect:
        add_label(doc, "Thấy gì là OK: ", expect, color=RGBColor(0x07, 0x89, 0xB0))
    if speak:
        add_body(doc, "Lời thoại (đọc khi đang quay):", bold=True, size=11)
        for para in speak.strip().split("\n\n"):
            add_body(doc, para.strip())
    if tip:
        add_label(doc, "Lưu ý: ", tip, color=RGBColor(0x16, 0xA3, 0x4A))


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("KỊCH BẢN DEMO — SỐ LIỆU LIVE (k6)")
    set_run_font(r, size=15, bold=True, color=RGBColor(0x1E, 0x3A, 0x8A))

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run(
        "Đề tài: Tối ưu hóa cold-start AI Serverless bằng Weight-Sharing (tmpfs)"
    )
    set_run_font(r2, size=11, bold=True)

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run(
        "Image v5  ·  Baseline MODEL_IO_MBPS=2  ·  Quay ~6–10 phút  ·  Chỉ đọc số LIVE"
    )
    set_run_font(r3, size=11, italic=True, color=RGBColor(0x47, 0x55, 0x69))

    add_body(
        doc,
        "Demo đo P99 cold-start bằng k6 (COLD_VUS=3) sau scale-to-zero — "
        "không dùng một lệnh curl đơn, không mở results/FINAL-REPORT.md.",
        bold=True,
        size=11,
    )
    add_body(
        doc,
        f"Kỳ vọng tham chiếu (dry-run ĐẠT gần nhất): Disk ~{REF_BASE_P99}, "
        f"RAM ~{REF_OPT_P99}, giảm ~{REF_IMPROVE}. "
        "Khi quay chỉ đọc đúng số script in ra trên terminal.",
        size=11,
    )

    add_heading_custom(doc, "Video chứng minh gì", level=1)
    add_bullet(doc, "Hai Knative Service cùng image v5, min-scale=0: Baseline (Disk) vs Optimized (RAM tmpfs).")
    add_bullet(
        doc,
        "Baseline có MODEL_IO_MBPS=2: mô phỏng storage bị giới hạn băng thông "
        "(HDD/NFS/cloud volume). Trên Minikube/Docker Desktop đĩa ảo ~400MB/s nên "
        "không lộ bottleneck I/O nếu không giới hạn.",
    )
    add_bullet(doc, "Optimized đọc thẳng tmpfs, không throttle.")
    add_bullet(doc, "Script demo: scale-to-zero → k6 3 VU → in LIVE P99 Disk vs RAM.")
    add_bullet(doc, "Kết luận chỉ dựa trên dòng ĐẠT / bảng LIVE P99 vừa chạy.")

    add_heading_custom(doc, "Bản đồ Terminal", level=1)
    rows = [
        ("Terminal 1", "Port-forward Kourier", "Giữ nguyên suốt demo"),
        ("Terminal 2", "Demo chính", "demo-cold-compare.sh — zoom font lớn"),
        ("Terminal 3", "Pod (tuỳ chọn)", "kubectl get pods — thấy 0 rồi scale lên"),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Cửa sổ", "Vai trò", "Quy tắc"]):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=True)
    for i, (a, b, c) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        table.rows[i].cells[2].text = c
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)

    # ===== A =====
    add_heading_custom(doc, "PHẦN A — CHUẨN BỊ (KHÔNG QUAY)", level=1)

    add_step(
        doc, "A1", "Docker Desktop",
        "Bật Docker",
        expect="Docker đang chạy trên menu bar.",
    )

    add_step(
        doc, "A2", "Terminal 2",
        "Vào thư mục project",
        command=(
            'cd "/Users/huyngoquang/Documents/Tài liệu thạc sĩ/'
            'Hệ tính toán phân bố nâng cao - NT2204.CH201/yolo-serverless"'
        ),
        expect="ls thấy test.jpg, loadtest-cold.js, scripts/demo-cold-compare.sh.",
    )

    add_step(
        doc, "A3", "Terminal 2",
        "Minikube + 2 service (image v5)",
        command=(
            "minikube status\n"
            "# nếu chưa Running: minikube start --driver=docker\n"
            "kubectl get ksvc\n"
            "kubectl get ksvc yolo-inference-baseline -o yaml | grep -E 'image:|MODEL_IO_MBPS' -A1\n"
            "command -v k6 || brew install k6"
        ),
        expect=(
            "ksvc READY; image …:v5; MODEL_IO_MBPS value \"2\"; lệnh k6 có sẵn."
        ),
        tip=(
            "Chưa có / sai image → bash scripts/build-image.sh && "
            "kubectl apply -f service.yaml && kubectl apply -f service-baseline.yaml"
        ),
    )

    add_step(
        doc, "A4", "Terminal 2",
        "Weights + xác nhận tmpfs",
        command=(
            "bash scripts/setup-weights.sh\n"
            "minikube ssh \"mountpoint /mnt/shared-weights; "
            "df -h /mnt/shared-weights /mnt/disk-weights; "
            "ls -lh /mnt/shared-weights/yolov8l-world.pt /mnt/disk-weights/yolov8l-world.pt\""
        ),
        expect=(
            "shared-weights is a mountpoint (tmpfs); "
            "disk-weights trên overlay/đĩa node; hai file .pt ~90M."
        ),
    )

    add_step(
        doc, "A5", "Terminal 1",
        "Port-forward Kourier (giữ nguyên)",
        command="kubectl port-forward -n kourier-system svc/kourier 8080:80",
        expect="Forwarding from 127.0.0.1:8080 -> 80",
    )

    add_step(
        doc, "A6", "Terminal 2",
        "Health check hai mode (bắt buộc thấy model_io_mbps)",
        command=(
            f"curl -s {URL_OPT}/health\n"
            f"curl -s {URL_BASE}/health"
        ),
        expect=(
            'optimized: "model_io_mbps":0.0 ; '
            'baseline: "model_io_mbps":2.0 ; cả hai status ok.'
        ),
        tip=(
            "Baseline thiếu model_io_mbps hoặc =0 → revision cũ. "
            "Apply lại service-baseline.yaml, chờ READY, curl lại."
        ),
    )

    add_step(
        doc, "A7", "Terminal 2",
        "Dry-run trước khi Record (BẮT BUỘC)",
        command="COLD_VUS=3 DROP_CACHE=1 bash scripts/demo-cold-compare.sh",
        expect=(
            f"Cuối cùng có LIVE P99 Disk ~{REF_BASE_P99} / RAM ~{REF_OPT_P99} / "
            f"Giảm ~{REF_IMPROVE} và dòng ĐẠT: RAM nhanh hơn Disk. "
            "(Số lệch vài giây vẫn OK nếu Disk >> RAM.)"
        ),
        tip=(
            "Chạy mất ~4–8 phút (Baseline lâu vì throttle 2MB/s). "
            "CẢNH BÁO / exit ≠ 0 → đừng quay; kiểm tra A3–A6 rồi chạy lại."
        ),
    )

    add_step(
        doc, "A8", "Terminal 3 (tuỳ chọn)",
        "Theo dõi Pod",
        command=(
            "while true; do clear; "
            "kubectl get pods -l serving.knative.dev/service; sleep 2; done"
        ),
        tip="Có thể bỏ nếu rối — Terminal 2 đủ.",
    )

    add_body(
        doc,
        "Xong A7 ĐẠT → zoom font Terminal 2 lớn → Record → phần B. "
        "Không mở FINAL-REPORT.",
        bold=True,
    )

    # ===== B =====
    add_heading_custom(doc, "PHẦN B — QUAY (SỐ LIỆU LIVE)", level=1)
    add_body(
        doc,
        "B3 chạy ~4–8 phút: Baseline (throttle) thường ~50–70s P99; Optimized ~15–25s. "
        "Trong lúc chờ: nói lời thoại — đừng cắt, đừng mở báo cáo cũ.",
        italic=True,
        size=10,
    )

    rows_b = [
        ("0:00–0:40", "B1 Intro", "Nói — giải thích Disk vs RAM + MODEL_IO_MBPS"),
        ("0:40–1:10", "B2 Hai service", "kubectl get ksvc"),
        ("1:10–…", "B3 Đo live k6", "COLD_VUS=3 bash scripts/demo-cold-compare.sh"),
        ("…–cuối", "B4 Đọc số + kết luận", "Bảng LIVE P99 trên Terminal 2"),
    ]
    table2 = doc.add_table(rows=1 + len(rows_b), cols=3)
    table2.style = "Table Grid"
    for i, h in enumerate(["Thời gian", "Bước", "Nhập vào đâu"]):
        table2.rows[0].cells[i].text = h
        for p in table2.rows[0].cells[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=True)
    for i, (a, b, c) in enumerate(rows_b, start=1):
        table2.rows[i].cells[0].text = a
        table2.rows[i].cells[1].text = b
        table2.rows[i].cells[2].text = c
        for cell in table2.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)

    add_step(
        doc, "B1", "Không gõ lệnh",
        "Intro — claim đề tài + vì sao có MODEL_IO_MBPS",
        record=True,
        speak=f"""Em xin demo: tối ưu hóa cold-start suy diễn AI trên Serverless bằng Weight-Sharing qua RAM tmpfs.

Có hai service: Baseline đọc weights từ đĩa; Optimized đọc từ RAM. Trên Minikube trong Docker Desktop, đĩa ảo rất nhanh nên bottleneck I/O không lộ. Baseline được cấu hình MODEL_IO_MBPS bằng 2 — mô phỏng storage bị giới hạn băng thông như HDD, NFS hoặc cloud volume trong thực tế. Optimized không giới hạn, đọc thẳng tmpfs.

Em sẽ scale-to-zero rồi dùng k6 gửi ba request cold đồng thời — trước Disk, sau RAM — và đọc P99 vừa đo trên màn hình. Kỳ vọng tham chiếu khoảng Disk {REF_BASE_P99}, RAM {REF_OPT_P99}, giảm khoảng {REF_IMPROVE}; số chính thức là số LIVE sắp hiện ra.""",
        tip="Không đọc FINAL-REPORT. Chưa gõ lệnh.",
    )

    add_step(
        doc, "B2", "Terminal 2",
        "Hai service RAM vs Disk",
        record=True,
        command="kubectl get ksvc",
        expect="yolo-inference và yolo-inference-baseline, READY True.",
        speak="""Hai service cùng image, min-scale bằng 0. Baseline gắn MODEL_IO_MBPS bằng 2 trên đường đọc đĩa; Optimized dùng tmpfs Weight-Sharing.""",
    )

    add_step(
        doc, "B3", "Terminal 2",
        "Chạy đối chứng LIVE (Baseline → Optimized)",
        record=True,
        command="COLD_VUS=3 DROP_CACHE=1 bash scripts/demo-cold-compare.sh",
        expect=(
            "Cuối cùng: LIVE P99 Disk … / RAM … / Giảm P99 …% "
            "và ĐẠT: RAM nhanh hơn Disk. "
            "Trong lúc chạy: scale-to-zero, drop cache (baseline), k6 summary từng bên. "
            "Baseline thường ~1 phút; Optimized nhanh hơn rõ."
        ),
        speak="""(Nói trong lúc script chạy — có thể lặp / kéo dài.)

Script đang ép Baseline về zero, xóa page cache, rồi k6 gửi ba request cold đồng thời. Baseline phải đọc weights với băng thông bị giới hạn nên P99 sẽ cao.

Sau đó làm tương tự với Optimized trên RAM tmpfs — không throttle.

P99 là độ trễ của request chậm nhất trong ba request cold đồng thời — metric khớp pipeline nghiệm thu đồ án.

Khi bảng LIVE P99 hiện ra, em sẽ đọc hai số và phần trăm giảm.""",
        tip=(
            "Không Ctrl+C. Không mở FINAL-REPORT. "
            "CẢNH BÁO / exit ≠ 0 → dừng Record, sửa A3–A7 rồi quay lại."
        ),
    )

    add_step(
        doc, "B4", "Terminal 2 — nhìn khối LIVE P99 (không gõ thêm)",
        "Đọc số vừa đo + kết luận",
        record=True,
        speak="""Em đọc kết quả vừa đo: P99 Baseline Disk khoảng … giây; P99 Optimized RAM khoảng … giây; giảm khoảng … phần trăm.

Như vậy, trong điều kiện storage bị giới hạn băng thông, Weight-Sharing qua tmpfs giảm đáng kể chi phí cold-start so với đọc đĩa khi Serverless scale-from-zero dưới tải đồng thời.

Em xin hết phần demo.""",
        tip=(
            f"Điền đúng số trên màn hình (tham chiếu gần nhất: Disk {REF_BASE_P99}, "
            f"RAM {REF_OPT_P99}, giảm {REF_IMPROVE}). "
            "Không viện dẫn FINAL-REPORT. Dừng Record sau câu kết."
        ),
    )

    # ===== C =====
    add_heading_custom(doc, "PHẦN C — CHEAT SHEET", level=1)
    add_body(doc, "Khi quay (Terminal 2):", bold=True)
    add_code(
        doc,
        "1) kubectl get ksvc\n"
        "2) COLD_VUS=3 DROP_CACHE=1 bash scripts/demo-cold-compare.sh\n"
        "3) Đọc khối LIVE P99 → kết luận → dừng Record",
    )
    add_body(doc, "Terminal 1:", bold=True)
    add_code(doc, "kubectl port-forward -n kourier-system svc/kourier 8080:80")

    add_heading_custom(doc, "Giải thích nhanh nếu thầy hỏi", level=1)
    add_bullet(
        doc,
        "Vì sao cần MODEL_IO_MBPS? Đĩa ảo Docker Desktop ~400MB/s; đọc 90MB ~0.2s, "
        "trong khi cold-start end-to-end ~7–20s → không thấy Disk vs RAM. "
        "2MB/s mô phỏng storage thực tế chậm hơn.",
    )
    add_bullet(
        doc,
        "Có gian lận không? Optimized không throttle; chỉ Baseline bị giới hạn I/O "
        "để đối chứng công bằng với giả thiết storage chậm. tmpfs vẫn là cơ chế Weight-Sharing.",
    )
    add_bullet(
        doc,
        "Vì sao không một curl? Một request dễ “may mắn”; k6 3 VU mới phản ánh "
        "cold-start dưới tải đồng thời sau scale-to-zero.",
    )

    add_heading_custom(doc, "Nếu sự cố", level=1)
    add_bullet(doc, "k6 not found → brew install k6")
    add_bullet(doc, "connection refused → kiểm tra Terminal 1 port-forward")
    add_bullet(doc, "health baseline thiếu model_io_mbps:2 → image/revision cũ; apply lại service-baseline.yaml")
    add_bullet(doc, "Baseline P99 < 25s hoặc RAM ≥ Disk → CẢNH BÁO; kiểm tra MODEL_IO_MBPS + tmpfs (A3–A7)")
    add_bullet(doc, "shared-weights không mountpoint → bash scripts/setup-weights.sh")
    add_bullet(doc, "Không cần FINAL-REPORT.md; không cần venv")

    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(16)
    fr = foot.add_run(
        "Sinh từ generate_demo_script.py — image v5 · MODEL_IO_MBPS=2 · live k6 · không dùng FINAL-REPORT."
    )
    set_run_font(fr, size=9, italic=True, color=RGBColor(0x64, 0x74, 0x8B))

    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    build()
